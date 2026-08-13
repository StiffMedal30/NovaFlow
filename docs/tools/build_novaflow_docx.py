"""Build the NovaFlow Word documentation from the markdown docs.

This is intentionally small and dependency-light. It supports the markdown
constructs used by the project docs: headings, paragraphs, bullets, numbered
lists, fenced code blocks, and simple pipe tables.
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "docs" / "novaflow-documentation.docx"

SOURCE_FILES = [
    ROOT / "docs" / "README.md",
    ROOT / "docs" / "architecture.md",
    ROOT / "docs" / "local-development.md",
    ROOT / "docs" / "configuration.md",
    ROOT / "docs" / "deployment.md",
    ROOT / "docs" / "operations.md",
    ROOT / "docs" / "testing.md",
    ROOT / "docs" / "aws-resource-inventory.md",
    ROOT / "docs" / "aws-recap-and-teardown.md",
    ROOT / "builder" / "ecs" / "README.md",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_run_font(run, font_name: str = "Calibri", size: int | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)


def add_code_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="CodeBlock")
    run = p.add_run(text if text else " ")
    set_run_font(run, "Consolas", 9)


def split_table_row(line: str) -> list[str]:
    text = line.strip().strip("|")
    return [cell.strip().replace("<br/>", "\n") for cell in text.split("|")]


def is_separator_row(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*", line))


def add_table(doc: Document, lines: list[str]) -> None:
    rows = [split_table_row(line) for line in lines if not is_separator_row(line)]
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "NovaFlow Table"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    set_table_width(table, 9360)

    col_width = max(1200, int(9360 / col_count))
    for row_idx, row in enumerate(rows):
        cells = table.rows[row_idx].cells
        for col_idx in range(col_count):
            value = row[col_idx] if col_idx < len(row) else ""
            cell = cells[col_idx]
            set_cell_width(cell, col_width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, "Calibri", 9)
            if row_idx == 0:
                run.bold = True
                set_cell_shading(cell, "E8EEF5")

    doc.add_paragraph()


def add_inline_code_aware_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas", 9)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, "Calibri", 11)


def add_static_contents(doc: Document) -> None:
    doc.add_heading("Contents", level=1)
    for source in SOURCE_FILES:
        title = first_heading(source)
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(title)


def first_heading(path: Path) -> str:
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            return line[2:].strip()
    return path.stem


def render_markdown_file(doc: Document, path: Path, first_section: bool) -> None:
    if not first_section:
        doc.add_page_break()

    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_language = ""
    table_buffer: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if table_buffer and (not stripped.startswith("|") or i == len(lines) - 1):
            if i == len(lines) - 1 and stripped.startswith("|"):
                table_buffer.append(line)
                i += 1
            add_table(doc, table_buffer)
            table_buffer = []
            continue

        if in_code:
            if stripped.startswith("```"):
                in_code = False
                code_language = ""
            else:
                add_code_paragraph(doc, line)
            i += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            code_language = stripped[3:].strip()
            if code_language:
                p = doc.add_paragraph()
                run = p.add_run(f"Code: {code_language}")
                run.bold = True
                run.font.color.rgb = RGBColor(31, 77, 120)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            table_buffer.append(line)
            i += 1
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 3)
            title = heading.group(2)
            doc.add_heading(title, level=level)
            i += 1
            continue

        if stripped.startswith("- "):
            add_inline_code_aware_paragraph(doc, stripped[2:], "List Bullet")
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            add_inline_code_aware_paragraph(doc, numbered.group(1), "List Number")
            i += 1
            continue

        add_inline_code_aware_paragraph(doc, stripped)
        i += 1

    if table_buffer:
        add_table(doc, table_buffer)


def create_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    if "CodeBlock" not in styles:
        code = styles.add_style("CodeBlock", 1)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code.font.size = Pt(9)
        code.paragraph_format.space_after = Pt(0)
        code.paragraph_format.left_indent = Inches(0.15)

    if "NovaFlow Table" not in styles:
        table_style = styles.add_style("NovaFlow Table", 3)
        table_style.font.name = "Calibri"
        table_style.font.size = Pt(9)


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("NovaFlow Documentation")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    create_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("NovaFlow Documentation")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(31, 58, 95)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "Architecture, local development, configuration, deployment, operations, testing, and AWS runbooks"
    )
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(89, 89, 89)

    generated = doc.add_paragraph()
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated_run = generated.add_run(f"Generated {date.today().isoformat()}")
    generated_run.font.size = Pt(9)
    generated_run.font.color.rgb = RGBColor(127, 127, 127)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_static_contents(doc)

    for index, source in enumerate(SOURCE_FILES):
        render_markdown_file(doc, source, first_section=False)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
